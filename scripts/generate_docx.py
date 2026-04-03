import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def generate_report():
    doc = Document()

    # --- Global Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set default paragraph format
    doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.styles['Normal'].paragraph_format.space_after = Pt(12)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.0

    # --- Margins ---
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Helper for Chapter Headings
    def add_chapter(text):
        p = doc.add_heading(text.upper(), level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = None
        doc.add_paragraph()

    # Helper for Section Headings
    def add_section(text):
        p = doc.add_heading(text.upper(), level=1)
        run = p.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = None
        doc.add_paragraph()

    # Load content
    md_path = 'Project_Report_VisionTalk_DDU.md'
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    active_table = None
    num_cols = 0

    for line in lines:
        line = line.strip()
        if not line:
            in_table = False
            active_table = None
            continue
        
        # Simple Markdown parsing
        if line.startswith('---'):
            doc.add_page_break()
            in_table = False
            active_table = None
            continue
            
        if line.startswith('# 🔴 CHAPTER'):
            title = line.split('CHAPTER')[1].strip(': ')
            add_chapter(title)
            in_table = False
            active_table = None
            continue
            
        elif line.startswith('# 🔴 SECTION 1:'):
            title = line.split('SECTION 1:')[1].strip()
            add_chapter(title)
            in_table = False
            active_table = None
            continue
            
        elif line.startswith('### '):
            title = line.replace('### ', '').strip()
            add_section(title)
            in_table = False
            active_table = None
            continue
            
        # Table Handling
        if line.startswith('|') and '|' in line:
            if '---' in line: # Skip separator row
                continue
            
            # Extract parts, ignoring leading/trailing empty strings from splits
            parts = [p.strip() for p in line.split('|')]
            if line.startswith('|'): parts = parts[1:]
            if line.endswith('|'): parts = parts[:-1]
            
            if not in_table:
                num_cols = len(parts)
                active_table = doc.add_table(rows=0, cols=num_cols)
                active_table.style = 'Table Grid'
                in_table = True
            
            row_cells = active_table.add_row().cells
            # Match lengths to avoid IndexError
            for i in range(min(len(parts), num_cols)):
                row_cells[i].text = parts[i]
            continue
        else:
            in_table = False
            active_table = None

        # Content Handling
        if line.startswith('*   **') or line.startswith('- **') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = re.sub(r'\*\*', '', line.strip('* -')).strip()
            p.text = text
        elif line.startswith('<center>') or line.endswith('</center>'):
            text = line.replace('<center>', '').replace('</center>', '').strip()
            if not text: continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.replace('**', ''))
            if '**' in text:
                run.bold = True
        else:
            text = re.sub(r'\*\*', '', line)
            p = doc.add_paragraph(text)

    # Save
    output_path = 'VisionTalk_Project_Report_DDU.docx'
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    generate_report()
